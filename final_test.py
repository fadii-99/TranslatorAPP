# import http.client
# import xml.etree.ElementTree as ET
# from langchain import LLMChain, PromptTemplate
# from langchain.llms import OpenAI

# def convert_docx_to_xliff(api_key, docx_file_path):
#     """
#     Convert a DOCX file to XLIFF using the provided API.
#     (Note: The sample API call does not actually upload the DOCX file,
#     so docx_file_path is not used here. In a real scenario, you may need
#     to adjust this code to include the file content.)
#     """
#     conn = http.client.HTTPSConnection("translated-matecat-filters-v1.p.rapidapi.com")
#     # For Arabic, change targetLocale to Arabic locale (e.g. ar-SA)
#     payload = "sourceLocale=en-GB&targetLocale=ar-SA"
#     headers = {
#         'x-rapidapi-key': api_key,
#         'x-rapidapi-host': "translated-matecat-filters-v1.p.rapidapi.com",
#         'Content-Type': "application/x-www-form-urlencoded"
#     }
#     conn.request("POST", "/api/v2/original2xliff", payload, headers)
#     res = conn.getresponse()
#     data = res.read()
#     xliff_str = data.decode("utf-8")
#     return xliff_str

# def save_xliff(xliff_str, output_file_path):
#     """Save the given XLIFF string to a file."""
#     with open(output_file_path, "w", encoding="utf-8") as f:
#         f.write(xliff_str)

# def find_parent(root, child):
#     """Helper to find the parent element of a given child in the XML tree."""
#     for parent in root.iter():
#         if child in list(parent):
#             return parent
#     return None

# def translate_trans_unit(trans_unit_str, llm_chain):
#     """
#     Use the LLM chain to translate the text content within a <trans-unit> block.
#     The prompt instructs the model to only translate text content and to leave all XML tags and attributes intact.
#     """
#     result = llm_chain.run(input=trans_unit_str)
#     return result

# def process_xliff_file(input_file_path, output_file_path, llm_chain):
#     """
#     Parse the input XLIFF file, process each <trans-unit> element using the translation chain,
#     replace the original elements with the translated ones, and save the new XML to output_file_path.
#     """
#     tree = ET.parse(input_file_path)
#     root = tree.getroot()
    
#     # Find all <trans-unit> elements (works even if there is a namespace)
#     for trans_unit in root.findall(".//trans-unit"):
#         # Convert the trans-unit element to a string
#         trans_unit_str = ET.tostring(trans_unit, encoding="unicode")
#         # Translate the trans-unit using the LLM chain (which preserves XML tags)
#         translated_trans_unit_str = translate_trans_unit(trans_unit_str, llm_chain)
#         # Try to parse the translated XML snippet back into an element
#         try:
#             new_element = ET.fromstring(translated_trans_unit_str)
#         except Exception as e:
#             print(f"Error parsing translated XML: {e}")
#             continue
#         # Find the parent of the trans-unit element and replace it with the translated element
#         parent = find_parent(root, trans_unit)
#         if parent is not None:
#             index = list(parent).index(trans_unit)
#             parent.remove(trans_unit)
#             parent.insert(index, new_element)
#         else:
#             print("Warning: Could not find parent for a trans-unit element.")
    
#     # Save the updated XML tree to the output file with an XML declaration
#     tree.write(output_file_path, encoding="utf-8", xml_declaration=True)

# def main():
#     # ----- Step 1: Convert DOCX to XLIFF -----
#     # Replace with your actual RapidAPI key.
#     rapidapi_key = "385d0a1310msh4e957f9d0387159p1c24e7jsn4e1c90b423c5"
#     # Path to your DOCX file (if needed, adjust conversion function to read file content)
#     docx_file_path = "input.docx"
#     intermediate_xliff_file = "intermediate.xliff"
#     translated_xliff_file = "translated.xliff"
    
#     # Convert the DOCX file to an XLIFF string
#     xliff_str = convert_docx_to_xliff(rapidapi_key, docx_file_path)
#     # Save the XLIFF string to an intermediate file
#     save_xliff(xliff_str, intermediate_xliff_file)
#     print(f"Intermediate XLIFF file saved as {intermediate_xliff_file}")

#     # ----- Step 2: Set up the LangChain translation chain -----
#     # Create a prompt template that instructs the LLM to translate only the text content,
#     # preserving the XML structure.
#     prompt_template = PromptTemplate(
#         input_variables=["input"],
#         template="""You are provided with an XML snippet containing a <trans-unit> element. Your task is to translate the English text within the XML to Arabic, but you must preserve the XML tags, attributes, and overall structure exactly as they are.

#         Specifically:
#         - Translate the text content inside the <source>, <seg-source>, and <target> tags, including the text inside nested tags such as <mrk> and <g>.
#         - Do NOT modify any XML tags or attributes.
#         - Ensure that only the text is translated and that the output retains the original XML structure and tag hierarchy.

#         Here is the input XML:
#         {input}

#         Output the translated XML snippet.
#         """
#             )
#     # Create the LLM instance (make sure your OpenAI API key is set in your environment)
#     llm = OpenAI(temperature=0)
#     llm_chain = LLMChain(llm=llm, prompt=prompt_template)
    
#     # ----- Step 3: Process the XLIFF file and translate each <trans-unit> -----
#     process_xliff_file(intermediate_xliff_file, translated_xliff_file, llm_chain)
#     print(f"Translated XLIFF file saved as {translated_xliff_file}")

# if __name__ == "__main__":
#     main()
# import http.client

# conn = http.client.HTTPSConnection("translated-matecat-filters-v1.p.rapidapi.com")

# payload = "-----011000010111000001101001\r\nContent-Disposition: form-data; name=\"document\"\r\n\r\ninput.docx\r\n-----011000010111000001101001\r\nContent-Disposition: form-data; name=\"sourceLocale\"\r\n\r\nen-GB\r\n-----011000010111000001101001\r\nContent-Disposition: form-data; name=\"targetLocale\"\r\n\r\nit-IT\r\n-----011000010111000001101001\r\nContent-Disposition: form-data; name=\"utf8FileName\"\r\n\r\ninput.docx\r\n-----011000010111000001101001--\r\n\r\n"

# headers = {
#     'x-rapidapi-key': "385d0a1310msh4e957f9d0387159p1c24e7jsn4e1c90b423c5",
#     'x-rapidapi-host': "translated-matecat-filters-v1.p.rapidapi.com",
#     'Content-Type': "multipart/form-data; boundary=---011000010111000001101001"
# }

# conn.request("POST", "/api/v2/original2xliff", payload, headers)

# res = conn.getresponse()
# data = res.read()

# print(data.decode("utf-8"))
import openai
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class DocxTranslator:
    def __init__(self, openai_api_key, target_language='Arabic'):
        # Set your OpenAI API key
        openai.api_key = openai_api_key
        self.client = openai
        self.target_language = target_language

    def translate_text(self, text):
        """Translate a piece of text using the OpenAI API."""
        try:
            response = self.client.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": f"You are a translator. Translate the following text to {self.target_language}."},
                    {"role": "user", "content": text}
                ],
                temperature=0.3
            )
            translated_text = response.choices[0].message.content.strip()
            return translated_text
        except Exception as e:
            print(f"Translation error: {e}")
            # In case of an error, return the original text to avoid data loss.
            return text

    def set_paragraph_rtl(self, paragraph):
        """
        Optionally set the paragraph direction to right-to-left.
        This is useful for Arabic layout.
        """
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def translate_paragraph(self, paragraph):
        """
        Translate each run in a paragraph.
        Note: Translating run-by-run works for many cases, but if your source text
        splits sentences across multiple runs, you might need to adjust the logic.
        """
        for run in paragraph.runs:
            if run.text.strip():
                run.text = self.translate_text(run.text)
        # Uncomment the next line to enforce right-to-left formatting on the paragraph:
        # self.set_paragraph_rtl(paragraph)

    def translate_table(self, table):
        """Recursively translate all paragraphs within a table (including nested tables)."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.translate_paragraph(paragraph)
                for nested_table in cell.tables:
                    self.translate_table(nested_table)

    def translate_document(self, input_path, output_path):
        """
        Load a DOCX file, translate its text using the OpenAI API, and save a new translated DOCX.
        This processes paragraphs, tables, headers, and footers.
        """
        doc = Document(input_path)

        # Translate main body paragraphs
        for paragraph in doc.paragraphs:
            self.translate_paragraph(paragraph)

        # Translate tables in the main document
        for table in doc.tables:
            self.translate_table(table)

        # Optionally, translate headers and footers (if present)
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self.translate_paragraph(paragraph)
            footer = section.footer
            for paragraph in footer.paragraphs:
                self.translate_paragraph(paragraph)

        doc.save(output_path)
        print(f"Translated document saved as '{output_path}'")

if __name__ == "__main__":
    # Replace with your actual OpenAI API key
    OPENAI_API_KEY = "sk-7DfCpAhhFAtLzUZGAe-sLWqSj6Vkg0uQ1-BkFg-SUUT3BlbkFJLATphGwsms69rMLCman_PyN7oYvnaJkyQtWrz1JzYA"
    translator = DocxTranslator(OPENAI_API_KEY, target_language="Arabic")
    
    # Define input and output file paths
    input_file = "input.docx"
    output_file = "translated.docx"
    
    translator.translate_document(input_file, output_file)
